"""
Document loader for various file formats
Handles PDF, DOCX, TXT, MD, and XLSX files
"""

import logging
from pathlib import Path
from typing import List, Optional
import os

# LlamaIndex document types
from llama_index.core import Document

logger = logging.getLogger(__name__)

class DocumentLoader:
    """Handles loading and processing of various document formats"""
    
    def __init__(self):
        self.supported_extensions = {'.txt', '.pdf', '.docx', '.md', '.xlsx'}
    
    def load_text_file(self, file_path: Path) -> Optional[Document]:
        """Load a text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return Document(
                text=content,
                metadata={
                    "filename": file_path.name,
                    "file_path": str(file_path),
                    "file_type": "text"
                }
            )
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {e}")
            return None
    
    def load_pdf_file(self, file_path: Path) -> Optional[Document]:
        """Load a PDF file"""
        try:
            import PyPDF2
            
            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            return Document(
                text=text,
                metadata={
                    "filename": file_path.name,
                    "file_path": str(file_path),
                    "file_type": "pdf"
                }
            )
        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            return None
        except Exception as e:
            logger.error(f"Error loading PDF file {file_path}: {e}")
            return None
    
    def load_docx_file(self, file_path: Path) -> Optional[Document]:
        """Load a DOCX file"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return Document(
                text=text,
                metadata={
                    "filename": file_path.name,
                    "file_path": str(file_path),
                    "file_type": "docx"
                }
            )
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return None
        except Exception as e:
            logger.error(f"Error loading DOCX file {file_path}: {e}")
            return None
    
    def load_excel_file(self, file_path: Path) -> Optional[Document]:
        """Load an Excel file"""
        try:
            import pandas as pd
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text_parts = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                text_parts.append(f"Sheet: {sheet_name}\n")
                text_parts.append(df.to_string())
                text_parts.append("\n\n")
            
            text = "".join(text_parts)
            
            return Document(
                text=text,
                metadata={
                    "filename": file_path.name,
                    "file_path": str(file_path),
                    "file_type": "excel"
                }
            )
        except ImportError:
            logger.error("pandas and openpyxl not installed. Install with: pip install pandas openpyxl")
            return None
        except Exception as e:
            logger.error(f"Error loading Excel file {file_path}: {e}")
            return None
    
    def load_single_document(self, file_path: Path) -> Optional[Document]:
        """Load a single document based on its extension"""
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None
        
        extension = file_path.suffix.lower()
        
        if extension == '.txt' or extension == '.md':
            return self.load_text_file(file_path)
        elif extension == '.pdf':
            return self.load_pdf_file(file_path)
        elif extension == '.docx':
            return self.load_docx_file(file_path)
        elif extension == '.xlsx':
            return self.load_excel_file(file_path)
        else:
            logger.warning(f"Unsupported file type: {extension}")
            return None
    
    def load_all_documents(self, directory: Path) -> List[Document]:
        """Load all supported documents from a directory"""
        documents = []
        
        if not directory.exists():
            logger.warning(f"Directory not found: {directory}")
            return documents
        
        logger.info(f"Loading documents from: {directory}")
        
        # Get all files in directory
        files = [f for f in directory.iterdir() if f.is_file()]
        
        for file_path in files:
            if file_path.suffix.lower() in self.supported_extensions:
                logger.info(f"Loading: {file_path.name}")
                doc = self.load_single_document(file_path)
                if doc:
                    documents.append(doc)
                else:
                    logger.warning(f"Failed to load: {file_path.name}")
            else:
                logger.info(f"Skipping unsupported file: {file_path.name}")
        
        logger.info(f"Successfully loaded {len(documents)} documents")
        return documents